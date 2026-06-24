"""Generate Word (.docx) newsletter document."""
import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1B3A5C"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_deal_row(doc, deal, emoji):
    val = deal.get("deal_value_usd_bn")
    val_str = f"${val:.1f}B" if val else "Undisclosed"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{emoji} {deal['headline']}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("1B3A5C")

    details = doc.add_paragraph()
    details.paragraph_format.left_indent = Inches(0.25)
    details.paragraph_format.space_after = Pt(1)
    r1 = details.add_run(f"Value: {val_str}  |  Status: {deal.get('status','—')}  |  Category: {deal.get('category','—')}  |  Geography: {deal.get('geography','—')}")
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor.from_string("555555")

    rationale = doc.add_paragraph()
    rationale.paragraph_format.left_indent = Inches(0.25)
    rationale.paragraph_format.space_after = Pt(6)
    r2 = rationale.add_run("Why it matters: ")
    r2.bold = True
    r2.font.size = Pt(9)
    r3 = rationale.add_run(deal.get("strategic_rationale", ""))
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor.from_string("333333")


def create_word(deals_path: str, output_path: str):
    with open(deals_path, encoding="utf-8") as f:
        deals = json.load(f)
    included = [d for d in deals if d.get("include_in_newsletter")]

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("🌍  FMCG M&A INTELLIGENCE NEWSLETTER")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("1B3A5C")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run(f"Issue Date: {datetime.now().strftime('%B %d, %Y')}  |  Powered by FMCG Intel Agent")
    rs.font.size = Pt(10)
    rs.font.color.rgb = RGBColor.from_string("888888")
    rs.italic = True

    doc.add_paragraph()

    # Executive Summary box (table)
    total_val = sum(d.get("deal_value_usd_bn", 0) or 0 for d in included)
    summary_table = doc.add_table(rows=1, cols=1)
    summary_table.style = "Table Grid"
    cell = summary_table.rows[0].cells[0]
    set_cell_bg(cell, "EAF2FB")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_head = p.add_run("EXECUTIVE SUMMARY\n")
    r_head.bold = True
    r_head.font.size = Pt(11)
    r_head.font.color.rgb = RGBColor.from_string("1B3A5C")
    r_body = p.add_run(
        f"Deals tracked: {len(included)}  |  Total disclosed value: ~${total_val:.1f}B  |  Period: 2025–2026 YTD\n\n"
        "The FMCG deal landscape has been defined by three forces:\n"
        "(1) Portfolio rationalisation — majors divesting non-core assets\n"
        "(2) New operating model acquisitions — buying DTC, functional & health brands\n"
        "(3) Category scale-up — mega-mergers in snacking, coffee and personal care"
    )
    r_body.font.size = Pt(10)

    doc.add_paragraph()

    # Section categoriser
    mega = [d for d in included if (d.get("deal_value_usd_bn") or 0) >= 10 and "Failed" not in str(d.get("status",""))]
    strategic = [d for d in included if 1 <= (d.get("deal_value_usd_bn") or 0) < 10]
    bolt = [d for d in included if not d.get("deal_value_usd_bn") and "Failed" not in str(d.get("status","")) and "PE" not in str(d.get("deal_type",""))]
    pe = [d for d in included if "PE" in str(d.get("deal_type","")) or "minority" in str(d.get("deal_type","")).lower()]
    failed = [d for d in included if "Failed" in str(d.get("status","")) or "Shelved" in str(d.get("status",""))]

    if mega:
        add_heading(doc, "SECTION 1: MEGA-DEALS (>$10B)", level=2)
        for d in mega:
            add_deal_row(doc, d, "🔷")

    if strategic:
        add_heading(doc, "SECTION 2: STRATEGIC ACQUISITIONS ($1B–$10B)", level=2)
        for d in strategic:
            add_deal_row(doc, d, "🔶")

    if bolt:
        add_heading(doc, "SECTION 3: BOLT-ON DEALS & UNDISCLOSED", level=2)
        for d in bolt:
            add_deal_row(doc, d, "🟡")

    if pe:
        add_heading(doc, "SECTION 4: PE / FUND ACTIVITY", level=2)
        for d in pe:
            add_deal_row(doc, d, "🏦")

    if failed:
        add_heading(doc, "SECTION 5: FAILED DEALS & DIVESTITURES", level=2)
        for d in failed:
            add_deal_row(doc, d, "❌")

    # Key Themes
    add_heading(doc, "KEY THEMES TO WATCH", level=2, color="0D7377")
    themes = [
        ("1. Health & Wellness Premium", "Functional beverages (Poppi, Alani Nu) command the highest valuation multiples (14x+ EBITDA); expect more bolt-ons in 2026."),
        ("2. DTC Operating Model Access", "Danone/Huel, Emami/ManCo signal FMCG majors buying digital-native capabilities and distribution, not just brands."),
        ("3. Portfolio Cleanup", "Unilever food exit, Nestlé water sale, Costa saga — conglomerates shedding operationally distinct categories."),
        ("4. GLP-1 Wildcard", "Weight-loss drug adoption reshaping snack/food M&A thesis; boards avoiding high-calorie category concentration."),
        ("5. EBITDA Multiples", "Food & bev settling at 10–11x; functional and wellness brands commanding 14x+; PE applying greater discipline."),
    ]
    for title_t, body_t in themes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{title_t}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(body_t)
        r2.font.size = Pt(10)

    # Pipeline Transparency
    doc.add_paragraph()
    add_heading(doc, "PIPELINE TRANSPARENCY", level=2, color="555555")
    transparency = doc.add_table(rows=6, cols=2)
    transparency.style = "Table Grid"
    rows_data = [
        ("Raw articles ingested", "16"),
        ("Duplicates removed", "2 (exact group match + headline similarity ≥0.75)"),
        ("After deduplication", "14 unique deals"),
        ("Credibility tiers", "T1=SEC/Official (9-10pt) | T2=Trade/Advisory (6-7pt) | T3=Blog (2-4pt)"),
        ("Relevance threshold", "Composite score ≥5.5 (60% relevance + 40% credibility)"),
        ("Assumptions", "Values at announcement; Undisclosed = no public figure found; Not investment advice"),
    ]
    for i, (k, v) in enumerate(rows_data):
        cell_k = transparency.rows[i].cells[0]
        cell_v = transparency.rows[i].cells[1]
        set_cell_bg(cell_k, "EAF2FB")
        cell_k.paragraphs[0].add_run(k).bold = True
        cell_k.paragraphs[0].runs[0].font.size = Pt(9)
        cell_v.paragraphs[0].add_run(v).font.size = Pt(9)
    transparency.columns[0].width = Cm(5)
    transparency.columns[1].width = Cm(12)

    # Footer
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer_p.add_run("FMCG Intel Agent  |  For internal use only  |  Not investment advice")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor.from_string("AAAAAA")
    rf.italic = True

    doc.save(output_path)
    print(f"[Word] Saved → {output_path}")


if __name__ == "__main__":
    create_word(
        "deals_final.json",
        "FMCG_MA_Newsletter.docx"
    )
